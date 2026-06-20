#!/usr/bin/env python3
"""
Génère les FICHES LOCAUX au format Word ÉDITABLE (.docx) — 1 page A4 paysage par
local, même gabarit que le PDF (en-tête + finitions/portes/CVC à gauche ;
électricité CFO/CFA + fluides + plomberie + observations à droite).

Source : docs/fiches-locaux/rooms-data.json (gen-fiches-locaux.mjs).
Sortie : docs/fiches-locaux/fiches-locaux-nassib.docx
"""
import os, json
from docx import Document
from docx.shared import Pt, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
DATA = json.load(open(os.path.join(BASE, "docs", "fiches-locaux", "rooms-data.json")))
OUT = os.path.join(BASE, "docs", "fiches-locaux", "fiches-locaux-nassib.docx")

INK = RGBColor(0x1C, 0x21, 0x28)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
CRIT = {"C": ("CRITIQUE", "DC2828"), "E": ("ÉLEVÉ", "EF8C14"),
        "M": ("MOYEN", "C9A21A"), "F": ("FAIBLE", "46AA5A")}
GROUPE = {"accueil_admin": "Groupe 1 : Administration", "urgences": "Groupe 2 : Urgences",
          "consultations": "Groupe 3 : Consultations", "gyneco_obstetrique": "Groupe 4 : Gynéco-Obstétrique",
          "bloc_cesarienne": "Groupe 5 : Bloc opératoire", "sspi": "Groupe 5 : Bloc opératoire",
          "sterilisation": "Groupe 6 : Stérilisation", "imagerie": "Groupe 7 : Imagerie",
          "laboratoire": "Groupe 8 : Laboratoire", "pharmacie": "Groupe 9 : Pharmacie",
          "hospitalisation": "Groupe 10 : Hospitalisation", "neonatologie": "Groupe 11 : Néonatologie",
          "vestiaires": "Groupe 12 : Locaux personnel", "locaux_techniques": "Groupe 13 : Technique / VRD"}


def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear"); sh.set(qn("w:fill"), hexcolor)
    tcPr.append(sh)


def set_borders(table, sz=4, color="BFC7CE"):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single"); e.set(qn("w:sz"), str(sz))
        e.set(qn("w:space"), "0"); e.set(qn("w:color"), color)
        borders.append(e)
    tblPr.append(borders)


def no_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "none")
        borders.append(e)
    tblPr.append(borders)


def set_widths(table, widths_mm):
    table.autofit = False
    for row in table.rows:
        for i, w in enumerate(widths_mm):
            row.cells[i].width = Mm(w)


def run(cell, text, size=8, bold=False, color=INK, align=None):
    p = cell.paragraphs[0]
    if p.runs or p.text:
        p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(0); p.paragraph_format.space_before = Pt(0)
    if align: p.alignment = align
    r = p.add_run(str(text))
    r.font.size = Pt(size); r.font.bold = bold; r.font.color.rgb = color
    r.font.name = "Calibri"
    return p


def section_title(cell, text, color="1C2128"):
    p = cell.paragraphs[0] if not (cell.paragraphs[0].text) else cell.add_paragraph()
    p.paragraph_format.space_after = Pt(1); p.paragraph_format.space_before = Pt(3)
    r = p.add_run(" " + text)
    r.font.size = Pt(8.5); r.font.bold = True; r.font.color.rgb = WHITE
    shade(cell, color)  # not ideal (whole cell) -> use a 1-row table instead


def kv_table(container, rows, w_label, w_val):
    t = container.add_table(rows=0, cols=2)
    set_borders(t); set_widths(t, [w_label, w_val])
    for k, v in rows:
        c = t.add_row().cells
        run(c[0], k, 7.5, color=RGBColor(0x55, 0x60, 0x70))
        run(c[1], v, 7.5, bold=True)
    return t


def grid_table(container, header, rows, widths):
    t = container.add_table(rows=1, cols=len(header))
    set_borders(t); set_widths(t, widths)
    for i, h in enumerate(header):
        run(t.rows[0].cells[i], h, 7, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER if i else None)
        shade(t.rows[0].cells[i], "DDE7EE")
    for r in rows:
        c = t.add_row().cells
        for i, val in enumerate(r):
            run(c[i], val, 7.5, bold=(i > 0), align=WD_ALIGN_PARAGRAPH.CENTER if i else None)
    return t


def band(container, text, color):
    """Titre de section : table 1x1 colorée."""
    t = container.add_table(rows=1, cols=1)
    no_borders(t)
    cell = t.rows[0].cells[0]
    shade(cell, color)
    run(cell, text, 8.5, bold=True, color=WHITE)
    cell.paragraphs[0].paragraph_format.space_after = Pt(0)
    return t


def fiche(doc, r, first):
    sec = doc.add_section() if not first else doc.sections[0]
    sec.orientation = WD_ORIENT.LANDSCAPE
    sec.page_width, sec.page_height = Mm(297), Mm(210)
    for m in ("top", "bottom", "left", "right"):
        setattr(sec, f"{m}_margin", Mm(8))

    critlabel, critcol = CRIT[r["crit"]]
    hb = r["hb"]
    isBloc = r["template"] in ("bloc_cesarienne", "bloc_operatoire", "sspi_reveil")
    isLab = r["template"] in ("laboratoire", "sterilisation")

    # ── En-tête ──
    h = doc.add_table(rows=1, cols=3)
    set_borders(h, sz=8, color="1C2128"); set_widths(h, [95, 150, 36])
    c0, c1, c2 = h.rows[0].cells
    run(c0, "K'BIO — Polyclinique Nassib · FIOG (Djibouti)", 8, bold=True, color=RGBColor(0, 0x3F, 0x72))
    run(c0, GROUPE.get(r["zone"], "Groupe — à préciser"), 11, bold=True)
    shade(c0, "EEF3F7")
    run(c1, r["name"], 13, bold=True)
    run(c1, f"Code local : {r['code']}  ·  Niveau {r['level']}  ·  Plan {r['sheet']}  ·  {r['dept']}", 8.5)
    run(c2, critlabel, 11, bold=True, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
    shade(c2, critcol)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # ── Corps : 2 colonnes ──
    body = doc.add_table(rows=1, cols=2)
    no_borders(body); set_widths(body, [138, 143])
    L, R = body.rows[0].cells

    # Colonne gauche
    band(L, "Sols / Murs / Plafonds", "1C2128")
    kv_table(L, [
        ("Charge au sol", f"{r['chargeSol']} daN/m²"),
        ("Sol", r["floor"]), ("Murs", r["walls"]),
        ("Plafond", r["ceiling"]), ("Plinthe", r["skirt"]),
        ("UPEC", r["upec"]), ("H. sous faux-plafond", f"{r['ceilingH']:.2f} m"),
    ], 42, 94)
    band(L, "Portes", "1C2128")
    kv_table(L, [
        ("Vantail", f"{r['doorW']}×{r['doorH']} mm — {r['doorColor']}"),
        ("Huisserie", r["doorFrame"]),
        ("Coupe-feu", r["doorFire"]), ("Contrôle d'accès", r["accessCtrl"]),
    ], 42, 94)
    band(L, "Chauffage / Ventilation / Climatisation", "0F7B6C")
    kv_table(L, [
        ("Climatisation", f"{r['coolKw']} kW — {r['ctaZone']}" if r["coolKw"] else "—"),
        ("Renouvellement d'air", f"{r['ach']} vol/h — {r['norm']}"),
        ("Température (été)", f"{r['tempC']} °C  ·  Hiver 22 °C"),
        ("Hygrométrie / Surpression", f"{r['humid'] or '—'} %  /  {('+'+str(r['surpression'])+' Pa') if r['surpression'] else '—'}"),
        ("Filtration", r["filtration"]),
    ], 50, 86)

    # Colonne droite
    band(R, "Électricité — Courants forts & faibles", "003F72")
    elec = [
        ("PC 16A normale", str(r["pc16"]), "", "✓" if not hb else "", "✓" if hb else ""),
        ("PC ondulée / secourue", str(r["ondule"]), "", "", "✓" if hb else ""),
        ("PC 20A normale", str(r["pc20"]), "", "✓" if r["pc20"] else "", ""),
        ("Alim. spécifique dédiée", str(r["ded"]), "", "✓" if r["ded"] else "", ""),
        ("RJ45", str(r["rj45"]), "", "✓", "✓" if hb else ""),
        ("Appel infirmière", str(r["nurse"]), "✓" if hb else "", "", ""),
        ("Interphone / Visiophone", str(r["intercom"]), "", "✓" if r["intercom"] else "", ""),
        ("Vidéosurveillance", str(r["cctv"]), "", "✓" if r["cctv"] else "", ""),
        ("Contrôle d'accès", str(r["access"]), "", "✓" if r["access"] else "", ""),
        ("Prise TV", str(r["tv"]), "", "✓" if r["tv"] else "", ""),
    ]
    grid_table(R, ["Désignation", "Qté", "Bras", "Mur", "GTL"], elec, [62, 12, 12, 12, 16])
    kv_table(R, [("IP / BAES / Terre", f"{r['ip']} · {r['baes']} BAES · {r['earth']} terre · Wi-Fi {'oui' if r['wifi'] else 'non'}")], 42, 90)

    band(R, "Fluides médicaux", "00785A")
    flu = [
        ("Oxygène (O₂)", str(r["o2"]), "✓" if hb else "", "✓" if isBloc else "", "✓" if hb else ""),
        ("Vide / aspiration", str(r["vide"]), "✓" if hb else "", "✓" if isBloc else "", "✓" if hb else ""),
        ("Air médical 3,5 bar", str(r["air"]), "✓" if (hb and r["air"]) else "", "✓" if isBloc else "", ""),
        ("Protoxyde d'azote (N₂O)", r["n2oNote"] or "—", "", "", ""),
    ]
    grid_table(R, ["Gaz", "Qté", "Bras", "Mur", "GTL"], flu, [62, 12, 12, 12, 16])

    band(R, "Surface · Plomberie sanitaire", "8A5A00")
    kv_table(R, [
        ("Surface utile / Volume", f"{r['area']} m²  /  {r['volume']} m³"),
        ("Eau froide / ECS", f"{r['ef'] or '—'} / {r['ec'] or '—'}"),
        ("Eau traitée / osmosée", r["eauTraitee"] or "—"),
        ("Siphons sol / EU / EV", f"{r['siphons']} / {r['eu']} / {r['ev']}"),
    ], 50, 86)

    band(R, "Observations", "5B3A87")
    obs = [f"Criticité : {critlabel}. Régime électrique : {r['regime']}."]
    if hb: obs.append("Bandeau de lit (BTDL) : O₂/Vide/élec. secourue + appel malade intégrés à la tête de lit.")
    if isBloc: obs.append("Liaison équipotentielle + contrôleur permanent d'isolement (CPI) avec report d'alarme.")
    if r["n2oNote"]: obs.append(f"N₂O / AGSS : {r['n2oNote']}.")
    if isLab: obs.append("Eau traitée/adoucie + extraction renforcée ; alim. analyseurs/autoclave dédiée.")
    obs.append("Données : visible plan + FIOG ; valeurs CFO/CFA/CVC déduites — à valider par les BE.")
    ot = R.add_table(rows=0, cols=1); set_borders(ot); set_widths(ot, [136])
    for o in obs:
        run(ot.add_row().cells[0], "• " + o, 7.5)

    # pied de page
    foot = doc.add_paragraph()
    foot.paragraph_format.space_before = Pt(2)
    fr = foot.add_run(f"Fiche local {r['code']} · {r['name']} · Polyclinique Nassib (FIOG) — document de conception, ne vaut pas plan d'exécution.")
    fr.font.size = Pt(6.5); fr.font.color.rgb = RGBColor(0x8A, 0x97, 0xA3)


doc = Document()
# style par défaut compact
st = doc.styles["Normal"]; st.font.name = "Calibri"; st.font.size = Pt(8)
st.paragraph_format.space_after = Pt(0)

LVLORD = {"RDC": 0, "R+1": 1, "exterieur": 2}
rooms = sorted(DATA, key=lambda x: (LVLORD.get(x["level"], 9), x["code"]))
for i, r in enumerate(rooms):
    fiche(doc, r, first=(i == 0))

doc.save(OUT)
print(f"DOCX écrit : {OUT} · {len(rooms)} fiches (1 page/local)")
